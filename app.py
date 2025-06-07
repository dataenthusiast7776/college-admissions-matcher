import streamlit as st
import pandas as pd
import re
import string
import textwrap
from datetime import datetime, timedelta
from datetime import date, time
from difflib import get_close_matches
from collections import Counter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import io
import os
import smtplib
from email.message import EmailMessage
from fpdf import FPDF
from math import ceil
import spacy
from io import BytesIO
from docx import Document


# ——— Stopwords & Keyword Extraction ———
STOPWORDS = {
    "a","an","the","and","or","but","if","then","with","on","in",
    "at","by","for","of","to","from","is","are","was","were","it",
    "this","that","these","those","as","be","has","have","had","i",
    "you","he","she","we","they","them","his","her","my","your",
}

def extract_keywords(text):
    text = text.translate(str.maketrans("", "", string.punctuation))
    tokens = text.lower().split()
    return [t for t in tokens if t not in STOPWORDS]

# ——— Normalization Helpers ———
def normalize_ethnicity(ethnicity):
    if pd.isna(ethnicity):
        return "unknown"
    e = ethnicity.lower()
    if any(x in e for x in ["indian","south asian","asian"]):
        return "asian"
    if "white" in e or "caucasian" in e:
        return "white"
    if "black" in e or "african american" in e:
        return "black"
    if "hispanic" in e or "latino" in e or "latina" in e or "latinx" in e:
        return "hispanic"
    if "native american" in e or "indigenous" in e:
        return "native american"
    if "middle eastern" in e or "arab" in e:
        return "middle eastern"
    return "other"

def normalize_gender(gender):
    if pd.isna(gender):
        return "unknown"
    g = str(gender).strip().lower()
    if g in ['male','m']:
        return 'male'
    if g in ['female','f']:
        return 'female'
    return 'unknown'

def clean_acceptances(raw):
    if pd.isna(raw) or not raw.strip():
        return ""
    parts = re.split(r"[\n,]+", raw)
    bad_kw = {
        "club","volunteer","internship","hook","income","essay",
        "activity","award","reflection","summary","miscellaneous",
        "consideration","recommendation","research","grades"
    }
    school_kw = {
        "university","college","institute","state","academy","school",
        "tech","polytechnic","poly","mit","stanford","harvard","princeton","yale"
    }
    good = []
    for p in parts:
        pl = p.strip().lower()
        if not pl or any(b in pl for b in bad_kw):
            continue
        if any(s in pl for s in school_kw) or pl in {"ea","ed","rea","rd"} or len(pl.split())<=8:
            good.append(p.strip())
    joined = ", ".join(good)
    return joined if len(joined)<=250 else ""

def match_profiles(df, gpa, sat, act, eth, gen, ec_query, use_gpa=True):
    df['Eth_norm'] = df['Ethnicity'].apply(normalize_ethnicity)
    df['Gen_norm'] = df['Gender'].apply(normalize_gender)
    df['acc_clean'] = df['acceptances'].apply(clean_acceptances)
    d = df[df['acc_clean']!=""].copy()

    if eth!="No filter":
        d = d[d['Eth_norm']==eth.lower()]
    if gen!="No filter":
        d = d[d['Gen_norm']==gen.lower()]

    if use_gpa and gpa is not None:
        d = d[(d['GPA']>=gpa-0.05)&(d['GPA']<=gpa+0.05)]

    if sat is not None:
        d = d[d['SAT_Score'].apply(lambda x: abs(x-sat)<=30 if not pd.isna(x) else False)]
    if act is not None:
        d = d[d['ACT_Score'].apply(lambda x: abs(x-act)<=1 if not pd.isna(x) else False)]

    d['EC_matches'] = [[] for _ in range(len(d))]
    if ec_query.strip():
        keywords = extract_keywords(ec_query)
        if keywords:
            def check_row(ec_text):
                if pd.isna(ec_text):
                    return False, []
                ec_lower = ec_text.lower()
                hits = [kw for kw in keywords if kw in ec_lower]
                if len(keywords)>=2:
                    return (len(hits)>=2, hits)
                return (len(hits)>=1, hits)
            mask = d['parsed_ECs'].apply(lambda txt: check_row(txt)[0])
            d = d[mask].copy()
            d['EC_matches'] = d['parsed_ECs'].apply(lambda txt: check_row(txt)[1])

    return d[['url','GPA','SAT_Score','ACT_Score','Ethnicity','Gender','acc_clean','EC_matches']]

def filter_by_colleges(df, colleges_input):
    d = df[df['acc_clean'] != ""]
    acc_col = d['acc_clean'].str.lower()

    # Check for OR logic if " or " is present (case insensitive)
    if " or " in colleges_input.lower():
        cols = [c.strip().lower() for c in re.split(r"\s+or\s+", colleges_input, flags=re.IGNORECASE)]
        pattern = "|".join([re.escape(c) for c in cols])
        mask = acc_col.str.contains(pattern, na=False)
    else:
        # Default to AND logic using commas
        cols = [c.strip().lower() for c in colleges_input.split(",") if c.strip()]
        mask = acc_col.apply(lambda s: all(c in s for c in cols))

    return d[mask][['url', 'GPA', 'SAT_Score', 'ACT_Score', 'Ethnicity', 'Gender', 'acc_clean', 'parsed_ECs']]

def load_data():
    drive_url = "https://drive.google.com/uc?export=download&id=1nZtwYcUX_KraxOTAOLg6-ZvKZnKMNpSg"
    try:
        return pd.read_csv(drive_url)
    except Exception as e:
        st.warning(f"Could not load remote data from Google Drive, using local file. Error: {e}")
        return pd.read_csv("master_data.csv")
        
def display_results(res):
    if res.empty:
        st.warning("0 matches found.")
    else:
        st.success(f"Found {len(res)} matching profiles:")
        for _, r in res.iterrows():
            ec_hits = r.get('EC_matches', [])
            ec_line = f"<br><b>ECs in common:</b> {', '.join(ec_hits)}" if ec_hits else ""
            st.markdown(f"""
            <div style="font-size:14px; line-height:1.4; margin-bottom:8px;">
              • <a href="{r['url']}" target="_blank">{r['url']}</a><br>
              GPA: {r['GPA']:.2f} | SAT: {r['SAT_Score']} | ACT: {r['ACT_Score']}<br>
              Ethnicity: {r['Ethnicity']} | Gender: {r['Gender']}<br>
              Acceptances: {r['acc_clean']}{ec_line}
            </div>
            """, unsafe_allow_html=True)

# ——— New College List Wizard ———
def is_valid_email(email):
    return re.match(r"[^@]+@[^@]+\.[^@]+", email)

def normalize_residency(residency):
    if pd.isna(residency):
        return "other"
    r = residency.lower()
    if "domestic" in r:
        return "domestic"
    if "international" in r:
        return "international"
    return "other"

def fuzzy_match_major(user_major, majors_list, cutoff=0.6):
    if not user_major.strip():
        return None
    matches = get_close_matches(user_major.lower(), [m.lower() for m in majors_list], n=1, cutoff=cutoff)
    return matches[0] if matches else None

from collections import Counter

import smtplib
from email.message import EmailMessage
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from collections import Counter
import re
import pandas as pd

def college_list_wizard(df):
    st.markdown("### 🎓 College List Wizard")
    st.info("Provide your academic profile and we’ll email you a personalized list of colleges!")

    # Inputs
    gpa = st.text_input("Enter your GPA (0.0–4.0):")
    test_score = st.text_input("Enter SAT (400–1600) or ACT (1–36):")
    major = st.text_input("Intended Major (please spell out full major, e.g., 'Computer Science'):")
    ecs = st.text_area("Describe your Extracurriculars:")
    domestic = st.checkbox("Domestic student? (leave unchecked for International)")
    email = st.text_input("Enter your Email:")

    # Email validation
    if email and not is_valid_email(email):
        st.warning("Please enter a valid email address.")
        return

    # Parse GPA
    try:
        gpa_val = float(gpa)
    except:
        gpa_val = None

    # Parse test score
    sat_val = act_val = None
    if test_score.strip().isdigit():
        sc = int(test_score.strip())
        if 1 <= sc <= 36:
            act_val = sc
            sat_val = sc * 45
        elif 400 <= sc <= 1600:
            sat_val = sc

    # Match major
    def match_major(user_major, majors_list):
        user_major_lower = user_major.strip().lower()
        for m in majors_list:
            if user_major_lower in m.lower():
                return m
        return None

    majors_list = df['Major'].dropna().unique()
    matched_major = match_major(major, majors_list)

    if st.button("Match Me!", disabled=not is_valid_email(email)):
        df2 = df.copy()

        # Residency
        df2['Residency_norm'] = df2['Residency'].apply(normalize_residency)
        target_res = "domestic" if domestic else "international"
        df2 = df2[df2['Residency_norm'] == target_res]

        # GPA filter
        if gpa_val is not None:
            df2 = df2[(df2['GPA'] >= gpa_val - 0.1) & (df2['GPA'] <= gpa_val + 0.1)]

        # SAT/ACT filter
        def sat_act_match(row):
            sat_ok = sat_val is not None and not pd.isna(row['SAT_Score']) and abs(row['SAT_Score'] - sat_val) <= 30
            act_ok = act_val is not None and not pd.isna(row['ACT_Score']) and abs(row['ACT_Score'] - act_val) <= 1
            conv_ok = act_val is not None and not pd.isna(row['SAT_Score']) and abs(row['SAT_Score'] - act_val*45) <= 30
            return sat_ok or act_ok or conv_ok

        if sat_val or act_val:
            df2 = df2[df2.apply(sat_act_match, axis=1)]

        # Major
        if matched_major:
            df2 = df2[df2['Major'] == matched_major]

        # ECs
        ec_keys = extract_keywords(ecs)
        if ec_keys:
            df2 = df2[df2['parsed_ECs'].apply(lambda txt: any(kw in str(txt).lower() for kw in ec_keys))]

        # Extract clean college names
        def extract_clean_colleges(raw):
            if not isinstance(raw, str) or not raw.strip():
                return []
            parts = re.split(r"[\n,]+", raw)
            indicators = [
                "university", "college", "institute", "school",
                "academy", "tech", "polytechnic", "poly", "mit",
                "stanford", "harvard", "princeton", "yale"
            ]
            cleaned = []
            for p in parts:
                seg = p.strip()
                if not seg:
                    continue
                name = seg.split("(", 1)[0].strip()
                low = name.lower()
                if any(ind in low for ind in indicators):
                    cleaned.append(name[:100])
            return cleaned

        df2["cleaned_list"] = df2["acceptances"].apply(extract_clean_colleges)
        all_schools = [school for sub in df2["cleaned_list"] for school in sub]
        counts = Counter([s.lower() for s in all_schools])

        # Build PDF
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter

        logo_path = "assets/logo.png"
        logo_width = 100  # adjust width as needed
        logo_height = 100  # adjust height as needed

        # Draw logo at top right corner
        # x position: page width - right margin - logo width
        # y position: page height - top margin - logo height
        right_margin = 40
        top_margin = 50

        c.drawImage(logo_path, width - right_margin - logo_width, height - top_margin - logo_height, 
            width=logo_width, height=logo_height, mask='auto')

        # Title and Timestamp
        c.setFont("Helvetica-Bold", 16)
        c.drawString(40, height - 50, "MatchMyApp - Personalized College List")
        c.setFont("Helvetica", 10)
        c.drawString(40, height - 70, "Generated on: " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

        y = height - 100

        # User inputs - including ECs
        c.setFont("Helvetica-Bold", 12)
        c.drawString(40, y, "📌 Your Inputs")
        y -= 20
        c.setFont("Helvetica", 10)

        user_inputs = [
            ("GPA", gpa),
            ("SAT", str(sat_val) if sat_val else "N/A"),
            ("ACT", str(act_val) if act_val else "N/A"),
            ("Major", major if major else "N/A"),
            ("Residency", "Domestic" if domestic else "International"),
            ("Extracurriculars", ecs if ecs.strip() else "N/A"),
            ("Email", email if email else "N/A"),
        ]

        for label, value in user_inputs:
            lines = textwrap.wrap(f"{label}: {value}", width=90)
            for line in lines:
                c.drawString(50, y, line)
                y -= 15

        # Space after inputs
        y -= 20
        
        # Matched Colleges (up to 20)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(40, y, "🎯 Matched Colleges")
        y -= 25
        
        c.setFont("Helvetica-Bold", 11)
        max_colleges = 10
        for school, cnt in counts.most_common(max_colleges):
            college_name = school.title()
            text = f"{college_name} — {cnt} acceptance(s)"
            c.setFillColorRGB(0, 0, 0.5)  # dark blue
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y, college_name)
        
            c.setFillColorRGB(0, 0, 0)
            c.setFont("Helvetica", 11)
            accept_text = f" — {cnt} acceptance(s)"
            width_name = c.stringWidth(college_name, "Helvetica-Bold", 12)
            c.drawString(50 + width_name, y, accept_text)
        
            url = next((r['url'] for _, r in df2.iterrows() if school in str(r['acceptances']).lower()), None)
            if url:
                y -= 15
                c.setFont("Helvetica-Oblique", 9)
                c.setFillColorRGB(0, 0, 1)  # blue for link
                c.drawString(60, y, "Reddit: " + url)
                c.linkURL(url, (60, y - 2, 60 + c.stringWidth("Reddit: " + url, "Helvetica-Oblique", 9), y + 10), relative=0)
        
            y -= 25
            c.setFillColorRGB(0, 0, 0)
        
            if y < 80:
                c.showPage()
                y = height - 50
        
        # Reddit Insights — only filtered posts, max 10
        y -= 10
        c.setFont("Helvetica-Bold", 12)
        c.drawString(40, y, "🧠 Profiles Like Yours")
        y -= 25
        c.setFont("Helvetica", 9)
        
        reddit_filtered = df2.head(10)  # assuming df2 already filtered
        
        if reddit_filtered.empty:
            c.drawString(50, y, "No Reddit posts available after applying your filters.")
            y -= 20
        else:
            for _, row in reddit_filtered.iterrows():
                url = row['url']
                gpa_post = row.get('GPA', None)
                sat_post = row.get('SAT_Score', None)
                act_post = row.get('ACT_Score', None)
                major_post = row.get('Major', None)
                text = f"GPA {gpa_post}, SAT {sat_post}, ACT {act_post}, Major {major_post}"
        
                c.setFillColorRGB(0, 0, 1)
                c.drawString(50, y, url)
                c.linkURL(url, (50, y - 2, 50 + c.stringWidth(url, "Helvetica", 9), y + 10), relative=0)
                y -= 12
        
                c.setFillColorRGB(0, 0, 0)
                for line in textwrap.wrap(text, width=110):
                    c.drawString(50, y, line)
                    y -= 12
        
                y -= 15
                if y < 80:
                    c.showPage()
                    y = height - 50
                    c.setFont("Helvetica", 9)
        
        # Notes Section with your exact text
        y -= 10
        c.setFont("Helvetica-Bold", 12)
        c.drawString(40, y, "💡 Notes")
        y -= 25
        c.setFont("Helvetica", 11)
        
        notes = """
        These colleges were suggested to you because past applicants with similar profiles and interests got into them. When building your college list, please make sure to consider a range of factors, including your class size preferences, location, campus culture, sports culture, and financial aid.
        
        Also, note that most of the top schools are committed to meeting your full demonstrated need, but do your research since there are a few exceptions! 
        
        If you found this helpful, do share our app with a friend to spread the joy of college application preparation!
        """
        
        for line in notes.strip().split('\n'):
            for wrapped_line in textwrap.wrap(line.strip(), width=110):
                c.drawString(50, y, wrapped_line)
                y -= 15
            y -= 5
        
        c.save()
        buffer.seek(0)



        # Email PDF
        try:
            msg = EmailMessage()
            msg["Subject"] = "Your MatchMyApp Personalized College List"
            msg["From"] = st.secrets["EMAIL_ADDRESS"]
            msg["To"] = email
            msg.set_content("Attached is your personalized list of colleges based on your inputs. Good luck!")

            msg.add_attachment(buffer.read(), maintype="application", subtype="pdf", filename="college_list.pdf")

            with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
                smtp.login(st.secrets["EMAIL_ADDRESS"], st.secrets["EMAIL_APP_PASSWORD"])
                smtp.send_message(msg)

            st.success("✅ PDF sent to your email! If it’s playing hide and seek, check your Spam folder just in case. ")
        except Exception as e:
            st.error(f"❌ Failed to send email: {e}")




def generate_and_render_timeline(num_early, num_rd, num_ed2, start_date, fafsa_eligible):
    timeline = []
    # Normalize start_date to datetime with time 00:00 if only date provided
    if isinstance(start_date, date) and not isinstance(start_date, datetime):
        current_date = datetime.combine(start_date, time.min)
    else:
        current_date = start_date

    # Helper to add tasks to timeline and increment date by 7 days
    def add_week(tasks):
        nonlocal current_date
        timeline.append({"week_start": current_date, "tasks": tasks})
        current_date += timedelta(days=7)

    # Week 1: Common App basics
    add_week([
        "Familiarize yourself with the Common App",
        "Fill out personal and parent information"
    ])

    # Early apps timeline
    early_weeks = 9  # weeks until Nov 1 from late Aug start
    if num_early > 0:
        essays_per_week = ceil(num_early / (early_weeks - 1))
        # Week 2: invite recommenders + start essays
        add_week([
            f"Work on approximately {essays_per_week} Early Action/Decision essays",
            "Invite recommenders"
        ])
        # Weeks 3 to 8: essay work
        for _ in range(early_weeks - 3):
            add_week([f"Work on approximately {essays_per_week} Early Action/Decision essays"])
        # Week 9: finalize early apps
        add_week(["Finalize and submit Early Action/Decision applications"])
    else:
        # If no early apps, skip these weeks
        current_date += timedelta(days=7 * early_weeks)

    # FAFSA week mid-November (Nov 15)
    if fafsa_eligible:
        fafsa_date = datetime(start_date.year, 11, 15)
        # Only add FAFSA if current_date hasn't passed Nov 15 yet
        if current_date <= fafsa_date:
            # Jump current_date to FAFSA date if needed
            if current_date < fafsa_date:
                current_date = fafsa_date
            add_week(["If eligible, fill out FAFSA and CSS Profile"])

    # Regular Decision timeline
    rd_weeks = 7  # Nov 8 to Jan 1 ish
    if num_rd > 0:
        essays_per_week = ceil(num_rd / rd_weeks)
        rd_start = datetime(start_date.year, 11, 8)
        current_date = max(current_date, rd_start)
        for _ in range(rd_weeks - 1):
            timeline.append({
                "week_start": current_date,
                "tasks": [f"Work on approximately {essays_per_week} Regular Decision essays"]
            })
            current_date += timedelta(days=7)
        timeline.append({
            "week_start": current_date,
            "tasks": ["Finalize and submit Regular Decision applications"]
        })
        current_date += timedelta(days=7)

    # ED2 timeline
    ed2_weeks = 3
    if num_ed2 > 0:
        essays_per_week = ceil(num_ed2 / ed2_weeks)
        ed2_start = datetime(start_date.year + 1, 1, 8)
        current_date = max(current_date, ed2_start)
        for _ in range(ed2_weeks - 1):
            timeline.append({
                "week_start": current_date,
                "tasks": [f"Work on approximately {essays_per_week} Early Decision 2 essays"]
            })
            current_date += timedelta(days=7)
        timeline.append({
            "week_start": current_date,
            "tasks": ["Finalize and submit Early Decision 2 applications"]
        })
        current_date += timedelta(days=7)

    # Final review week
    timeline.append({
        "week_start": current_date,
        "tasks": ["Final review and submission of any remaining application materials"]
    })

    # Sort timeline chronologically just in case
    timeline.sort(key=lambda x: x["week_start"])

# ----------- PDF generation -------------

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "College Application Timeline", ln=True, align="C")
    pdf.ln(10)
    pdf.set_font("Arial", size=12)

    current_month = ""
    for entry in timeline:
        month_name = entry["week_start"].strftime("%B %Y")
        if month_name != current_month:
            current_month = month_name
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, current_month, ln=True)
            pdf.set_font("Arial", size=12)
        pdf.cell(0, 8, entry["week_start"].strftime("%b %d"), ln=True)
        for task in entry["tasks"]:
            pdf.multi_cell(0, 8, f" - {task}")
        pdf.ln(2)

    pdf_bytes = pdf.output(dest='S').encode('latin1')
    pdf_output = io.BytesIO(pdf_bytes)
    pdf_output.seek(0)

    st.download_button(
        label="Download Timeline PDF",
        data=pdf_output,
        file_name="college_application_timeline.pdf",
        mime="application/pdf"
    )

    # Also display in streamlit in text for quick preview
    for entry in timeline:
        st.markdown(f"### {entry['week_start'].strftime('%b %d, %Y')}")
        for task in entry["tasks"]:
            st.markdown(f"- {task}")




import re
from collections import Counter

theme_advice = {
    "belonging": "Talk about communities or spaces where you feel most at home. Think culture, clubs, religion, identity — anything that gives you a sense of place.",
    "personal identity": "This is your moment to reflect on how your background, experiences, or quirks shape who you are.",
    "cultural background": "You could reflect on your family traditions, heritage, or values passed down — and how they influence your choices and goals.",
    "intellectual heritage": "Highlight the ideas, books, or people that shaped your thinking. How do they still echo in what you study or care about today?",
    "personal growth": "Think about key turning points, setbacks, or revelations that taught you something real. How have you evolved?",
    "institutional values": "Make sure you understand what the college values — and show how you naturally align with those values, even if indirectly.",
    "academic opportunities": "This is a great place to name-drop programs, research, professors, or unique classes — show you’ve done your homework!",
}

# Use a simple regex + keyword-based approach
def extract_verbs(text):
    verbs_guess = re.findall(r'\b\w+ing\b|\b(to )?\w+\b', text.lower())
    common_ignore = {"being", "having", "doing", "getting", "making", "going", "saying", "is", "are", "was", "were", "be", "have", "do", "get", "make", "go", "say"}
    filtered = [v for v in verbs_guess if v not in common_ignore and len(v) > 3]
    return list(set(filtered))[:5] or ["reflect", "discuss"]

def extract_themes(text):
    text = text.lower()
    keywords = [
        "community", "membership", "identity", "race", "heritage",
        "growth", "challenge", "university", "school", "program",
        "opportunity", "academics", "culture", "tradition"
    ]
    theme_map = {
        "community": "belonging",
        "membership": "belonging",
        "identity": "personal identity",
        "race": "cultural background",
        "heritage": "intellectual heritage",
        "growth": "personal growth",
        "challenge": "personal growth",
        "school": "institutional values",
        "university": "institutional values",
        "program": "academic opportunities",
        "opportunity": "academic opportunities",
    }

    matched = [theme_map[k] for k in keywords if k in text]
    return list(set(matched)) or ["personal reflection"]

def analyze_prompt_nlp(prompt_text, word_limit):
    cleaned_text = re.sub(r'\s+', ' ', prompt_text.strip())
    length = len(cleaned_text.split())
    length_desc = "short" if length < 30 else "detailed"

    verbs = extract_verbs(prompt_text)
    themes = extract_themes(prompt_text)
    theme_text = f"**{', '.join(themes)}**" if themes else "broad personal reflection"
    first_theme = themes[0] if themes else None

    research_keywords = ["why this school", "why our university", "program", "opportunity", "major"]
    is_research = any(word in prompt_text.lower() for word in research_keywords)

    output = f"""

**Themes & What to Say:**  
Looks like your prompt touches on: {theme_text}

- Talk about communities or spaces where you feel most at home. Think culture, clubs, religion, identity — anything that gives you a sense of place.  
- This is your moment to reflect on how your background, experiences, or quirks shape who you are.  

{f'**Start here:** What’s one story that connects you to *{first_theme}*?' if first_theme else ''}

**Reflection Tips:**  
Think about moments in your life that changed how you see yourself. What sparked growth or gave you clarity?

**Research Tips:**  
{"This prompt is more about *you* than the school. But make sure your story still fits what the college values." if not is_research else "This prompt mentions the school. Research specific programs, values, or professors to tie your story to."}

**Prompt Length:**  
It’s {length_desc} ({length} words), so aim for clarity and emotional punch.

**Verbs to Respond To:**  
The prompt is nudging you to *{', '.join(verbs)}*. Let that guide your structure.

**Word Count Strategy:**  
~{word_limit} words should give you enough space to be real, but stay focused.
"""

    return output.strip()


def create_docx(prompt_text, essay_text):
    doc = Document()
    doc.add_heading("Essay Prompt", level=1)
    doc.add_paragraph(prompt_text)

    doc.add_paragraph("")  # First line break

    doc.add_heading("Your Essay Draft", level=1)
    doc.add_paragraph(essay_text)

    # Save to in-memory bytes buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ——— Main App ———
def main():
    st.markdown("""
    <div style="background-color: #1ABC9C; padding: 20px; text-align: center; font-size: 1.5em; font-weight: bold; color: #000;">
      We are taking your feedback! Click <a href="https://docs.google.com/forms/d/e/1FAIpQLSdgaCUa7S2KFfs6hUsFyDtttUZYiT46uTWtXEhhR9in8fEy6g/viewform?usp=header" target="_blank" style="color: #6A0DAD; text-decoration: underline;">here</a> to fill out a quick survey!
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div style="display: flex; justify-content: space-between; align-items: center;">
          <div style="flex-grow:1; text-align:center;">
            <h1 style='color:#6A0DAD; font-size:3em; margin:0;'>MatchMyApp</h1>
            <p style='color:#DAA520; font-size:1.2em; margin:0; line-height:1; margin-top:-8px; transform: translateX(-10px);'>
              Find your college application twin!
            </p>
          </div>
        </div>
    
        <div style="margin-top: 30px; max-width: 700px; margin-left: auto; margin-right: auto;">
          <p>I got bored one day, so I wrote a script to mine data off of the r/collegeresults subreddit. Well, one thing turned into another, and I realized I had a treasure trove of data to be put to use. A few days of caffeine-induced coding later, voila! MatchMyApp was born.</p>
    
          <p>Whether you're a junior preparing for college applications, a data enthusiast, or a parent looking to see how far your child could go, MatchMyApp has free, data-driven tools for you! Input your stats and see the results of similar past applicants in seconds, or build a targeted college list based on past acceptance data! Or, if you're a data nerd like me, head over to the data corner for an endless array of interesting graphs made from the master dataset.</p>
    
          <p>MatchMyApp is a work-in-progress, and I am currently working on adding LLM-supported features such as essay revision and guidance, personalized advice for extracurriculars, and more, all for free!</p>
        </div>
    
        <!-- DataDorm Button Section -->
        <div style="text-align: center; margin-top: 20px;">
          <a href="https://datadorm.streamlit.app" target="_blank" style="
              background-color: #D32F2F;
              color: white;
              padding: 20px 40px;
              font-size: 2em;
              font-weight: bold;
              border-radius: 12px;
              text-decoration: none;
              display: inline-block;
              box-shadow: 0 4px 8px rgba(0,0,0,0.2);
              transition: background-color 0.3s ease;
          " onmouseover="this.style.backgroundColor='#B71C1C';" onmouseout="this.style.backgroundColor='#D32F2F';">
            🏫 DataDorm
          </a>
          <p style="margin-top: 12px; font-size: 1.1em; color: #add8e6;">
            Check out our college admissions data search engine, sourced from official Common Data Sets!
          </p>
        </div>
        """, unsafe_allow_html=True)



    df = load_data()
    st.markdown("""
    <style>
    /* For the tab labels */
    div[role="tab"] {
        font-size: 20px !important;
        font-weight: 700 !important;
    }
    </style>
""", unsafe_allow_html=True)
    
    tabs = st.tabs(["Profile Filter", "Filter by College Acceptances", "College List Wizard", "Essay Timeline Planner", "Prompt Shop - Breakdown Essay Prompts"])

    with tabs[0]:
        st.markdown("#### Enter your profile (leave filters blank to skip):")
        use_gpa = st.checkbox("Filter by GPA", value=True)
        if use_gpa:
            gpa_s = st.slider("GPA (max 4.0)", 0.0, 4.0, 4.0, 0.01)
            gpa_m = st.number_input("Or enter GPA manually", 0.0, 4.0, gpa_s, 0.01)
            user_gpa = gpa_m if gpa_m != gpa_s else gpa_s
        else:
            user_gpa = None

        score_choice = st.selectbox("Score filter", ["No filter","SAT","ACT"])
        user_sat = user_act = None
        if score_choice=="SAT":
            user_sat = st.number_input("SAT Score",400,1600,1580,10)
        elif score_choice=="ACT":
            user_act = st.number_input("ACT Score",1,36,35,1)

        user_eth = st.selectbox(
            "Ethnicity",
            ["No filter","Asian","White","Black","Hispanic","Native American","Middle Eastern","Other"],
        )
        user_gen = st.selectbox("Gender", ["No filter","Male","Female"])
        ec_query = st.text_area(
            "Describe your extracurriculars:",
            placeholder="e.g., robotics club, varsity soccer, volunteer tutoring",
            height=80,
        )

        res = match_profiles(
            df, user_gpa, user_sat, user_act,
            user_eth, user_gen, ec_query,
            use_gpa=use_gpa
        )
        display_results(res)

    with tabs[1]:
        st.markdown("#### Filter profiles accepted to the following college(s):")
        college_input = st.text_input("Enter college name(s), comma‑separated. Use keyword OR to get profiles that were accepted to at-least one of the chosen colleges!")
        if college_input.strip():
            res = filter_by_colleges(df, college_input)
            display_results(res)
        else:
            st.info("Enter one or more college names to see matching acceptances.")

    with tabs[2]:
        college_list_wizard(df)

    with tabs[3]:
        st.markdown("### College Application Timeline Planner")

        with st.form("timeline_form"):
            num_early = st.number_input("Number of Early (REA, EA, ED1) colleges", min_value=0, step=1)
            num_rd = st.number_input("Number of Regular Decision colleges", min_value=0, step=1)
            num_ed2 = st.number_input("Number of Early Decision 2 colleges", min_value=0, step=1)
            fafsa_eligible = st.checkbox("Eligible for FAFSA and CSS Profile")
            start_date = st.date_input("Start date", value=datetime(2025, 8, 21))
            submitted = st.form_submit_button("Generate Timeline")

        if submitted:
            generate_and_render_timeline(num_early, num_rd, num_ed2, start_date, fafsa_eligible)

    with tabs[4]:
        st.markdown("### Prompt Shop — Prompt Breakdown Analyzer")

        # Initialize state variables if not set
        if 'show_drafting' not in st.session_state:
            st.session_state['show_drafting'] = False
        if 'breakdown_text' not in st.session_state:
            st.session_state['breakdown_text'] = ''
        if 'prompt_text' not in st.session_state:
            st.session_state['prompt_text'] = ''

        with st.form("prompt_form"):
            prompt_text = st.text_area("Paste your essay prompt here", height=150)
            word_limit = st.number_input("Word count limit", min_value=50, max_value=1000, value=650, step=10)
            submit_prompt = st.form_submit_button("Analyze Prompt")

        # Handle form submission
        if submit_prompt:
            if not prompt_text.strip():
                st.warning("Please enter an essay prompt to analyze.")
            else:
                breakdown = analyze_prompt_nlp(prompt_text, word_limit)
                st.session_state['breakdown_text'] = breakdown
                st.session_state['prompt_text'] = prompt_text
                st.session_state['show_drafting'] = False  # reset in case it's already True

        # Show breakdown if available
        if st.session_state['breakdown_text']:
            st.markdown("#### Your Prompt Breakdown:")
            st.markdown(st.session_state['breakdown_text'], unsafe_allow_html=True)

            # Button to open drafting space — outside form
            if st.button("📝 Start Drafting"):
                st.session_state['show_drafting'] = True

        # Show drafting section if toggled on
        if st.session_state['show_drafting']:
            st.markdown("## ✍️ Drafting Space")
            col1, col2 = st.columns([2, 2])

            with col1:
                st.markdown("### Breakdown Reference")
                st.markdown(st.session_state['breakdown_text'], unsafe_allow_html=True)

            with col2:
                essay_text = st.text_area("Your Essay Draft", height=400, key="essay_draft_text")

                if st.button("Download Essay + Breakdown (.docx)"):
                    docx_buffer = create_docx(st.session_state['prompt_text'], essay_text)
                    st.download_button(
                        label="Open Prepared File!",
                        data=docx_buffer,
                        file_name="essay_with_breakdown.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                                    

if __name__ == "__main__":
    main()
